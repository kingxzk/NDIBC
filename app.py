#!/usr/bin/env python3
"""
网络设备巡检和配置备份软件 - 主程序
支持思科、华为、H3C、锐捷、等厂商
"""

import os
import json
import time
import zipfile
import tempfile
import threading
import io
import base64
import sys
import webbrowser
from datetime import datetime
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, session, send_from_directory
from flask_cors import CORS
import netmiko
from netmiko import ConnectHandler
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import difflib
import hashlib
import pandas as pd

app = Flask(__name__, 
           static_folder='static',
           template_folder='templates')
app.secret_key = 'network-device-management-secret-key-2024'
CORS(app)

# 配置文件存储
CONFIG_DIR = Path("config_backups")
INSPECTION_DIR = Path("inspection_reports")
TEMPLATE_DIR = Path("templates")
CONFIG_DIR.mkdir(exist_ok=True)
INSPECTION_DIR.mkdir(exist_ok=True)
TEMPLATE_DIR.mkdir(exist_ok=True)

# 设备厂商映射
DEVICE_TYPES = {
    'cisco': 'cisco_ios',
    'huawei': 'huawei',
    'h3c': 'hp_comware',
    'ruijie': 'ruijie_os',
    'dell': 'dell_os10',
    'juniper': 'juniper_junos',
    'arista': 'arista_eos'
}

# 默认巡检命令
DEFAULT_INSPECTION_COMMANDS = {
    'cisco': [
        ('show version', '设备版本信息'),
        ('show running-config', '运行配置'),
        ('show interfaces status', '接口状态'),
        ('show ip interface brief', 'IP接口摘要'),
        ('show cdp neighbors', 'CDP邻居'),
        ('show logging', '系统日志'),
        ('show processes cpu', 'CPU使用率'),
        ('show memory statistics', '内存使用率'),
        ('show environment', '环境状态')
    ],
    'huawei': [
        ('display version', '设备版本信息'),
        ('display current-configuration', '当前配置'),
        ('display interface brief', '接口摘要'),
        ('display ip interface brief', 'IP接口摘要'),
        ('display lldp neighbor', 'LLDP邻居'),
        ('display logbuffer', '日志缓冲区'),
        ('display cpu-usage', 'CPU使用率'),
        ('display memory-usage', '内存使用率'),
        ('display device', '设备信息')
    ],
    'h3c': [
        ('display version', '设备版本信息'),
        ('display current-configuration', '当前配置'),
        ('display interface brief', '接口摘要'),
        ('display ip interface brief', 'IP接口摘要'),
        ('display lldp neighbor-information', 'LLDP邻居信息'),
        ('display logbuffer', '日志缓冲区'),
        ('display cpu-usage', 'CPU使用率'),
        ('display memory', '内存使用率'),
        ('display device', '设备信息')
    ],
    'ruijie': [
        ('show version', '设备版本信息'),
        ('show running-config', '运行配置'),
        ('show interfaces status', '接口状态'),
        ('show ip interface brief', 'IP接口摘要'),
        ('show lldp neighbors', 'LLDP邻居'),
        ('show logging', '系统日志'),
        ('show processes cpu', 'CPU使用率'),
        ('show memory', '内存使用率')
    ],
    'dell': [
        ('show version', '设备版本信息'),
        ('show running-config', '运行配置'),
        ('show interfaces status', '接口状态'),
        ('show ip interface brief', 'IP接口摘要'),
        ('show lldp neighbors', 'LLDP邻居'),
        ('show logging', '系统日志'),
        ('show processes cpu', 'CPU使用率'),
        ('show memory', '内存使用率')
    ]
}

class DeviceManager:
    """设备管理类"""
    
    def __init__(self):
        self.devices = {}
        self.load_devices()
    
    def load_devices(self):
        """加载设备列表"""
        try:
            if os.path.exists('devices.json'):
                with open('devices.json', 'r', encoding='utf-8') as f:
                    self.devices = json.load(f)
        except:
            self.devices = {}
    
    def save_devices(self):
        """保存设备列表"""
        with open('devices.json', 'w', encoding='utf-8') as f:
            json.dump(self.devices, f, ensure_ascii=False, indent=2)
    
    def add_device(self, name, ip, vendor, username, password, port=22):
        """添加设备"""
        device_id = hashlib.md5(f"{ip}:{port}".encode()).hexdigest()[:8]
        self.devices[device_id] = {
            'id': device_id,
            'name': name,
            'ip': ip,
            'vendor': vendor,
            'username': username,
            'password': password,
            'port': port,
            'created_at': datetime.now().isoformat()
        }
        self.save_devices()
        return device_id
    
    def batch_import(self, excel_data):
        """批量导入设备"""
        results = {
            'success': 0,
            'failed': 0,
            'errors': []
        }
        
        try:
            # 读取Excel数据
            df = pd.read_excel(io.BytesIO(excel_data))
            
            # 检查必要列
            required_columns = ['设备名称', 'IP地址', '厂商', '用户名', '密码']
            for col in required_columns:
                if col not in df.columns:
                    raise ValueError(f"缺少必要列: {col}")
            
            # 处理每一行
            for index, row in df.iterrows():
                try:
                    name = str(row['设备名称']).strip()
                    ip = str(row['IP地址']).strip()
                    vendor = str(row['厂商']).strip().lower()
                    username = str(row['用户名']).strip()
                    password = str(row['密码']).strip()
                    port = int(row.get('端口', 22))
                    
                    # 验证数据
                    if not name or not ip or not vendor or not username or not password:
                        raise ValueError("必要字段不能为空")
                    
                    if vendor not in DEVICE_TYPES:
                        raise ValueError(f"不支持的厂商: {vendor}")
                    
                    # 添加设备
                    device_id = self.add_device(name, ip, vendor, username, password, port)
                    results['success'] += 1
                    
                except Exception as e:
                    results['failed'] += 1
                    results['errors'].append({
                        'row': index + 2,  # Excel行号（从2开始）
                        'error': str(e),
                        'data': row.to_dict()
                    })
            
            return True, results
            
        except Exception as e:
            return False, str(e)
    
    def remove_device(self, device_id):
        """删除设备"""
        if device_id in self.devices:
            del self.devices[device_id]
            self.save_devices()
            return True
        return False
    
    def get_device(self, device_id):
        """获取设备信息"""
        return self.devices.get(device_id)
    
    def get_all_devices(self):
        """获取所有设备"""
        return list(self.devices.values())

class NetworkDevice:
    """网络设备连接类"""
    
    def __init__(self, device_info):
        self.device_info = device_info
        self.connection = None
    
    def connect(self):
        """连接到设备"""
        try:
            device_params = {
                'device_type': DEVICE_TYPES.get(self.device_info['vendor'], 'cisco_ios'),
                'host': self.device_info['ip'],
                'username': self.device_info['username'],
                'password': self.device_info['password'],
                'port': self.device_info.get('port', 22),
                'timeout': 30,
                'secret': self.device_info.get('enable_password', ''),
                'verbose': False
            }
            
            self.connection = ConnectHandler(**device_params)
            return True, "连接成功"
        except Exception as e:
            return False, f"连接失败: {str(e)}"
    
    def disconnect(self):
        """断开连接"""
        if self.connection:
            self.connection.disconnect()
            self.connection = None
    
    def execute_command(self, command):
        """执行命令"""
        try:
            if not self.connection:
                success, message = self.connect()
                if not success:
                    return False, message
            
            output = self.connection.send_command(command, delay_factor=2)
            return True, output
        except Exception as e:
            return False, f"执行命令失败: {str(e)}"
    
    def backup_config(self):
        """备份配置"""
        try:
            if not self.connection:
                success, message = self.connect()
                if not success:
                    return False, message
            
            # 获取配置
            if self.device_info['vendor'] in ['huawei', 'h3c']:
                config = self.connection.send_command('display current-configuration')
            else:
                config = self.connection.send_command('show running-config')
            
            # 保存配置
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{self.device_info['ip']}_{timestamp}.cfg"
            filepath = CONFIG_DIR / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(config)
            
            return True, {
                'filename': filename,
                'filepath': str(filepath),
                'config': config[:500] + "..." if len(config) > 500 else config
            }
        except Exception as e:
            return False, f"备份配置失败: {str(e)}"

def create_inspection_report(device_info, inspection_results, output_path):
    """创建巡检报告Word文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    
    # 标题 - 安全地创建标题
    try:
        title = doc.add_heading('网络设备巡检报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:  # 检查是否有runs
            title_run = title.runs[0]
            title_run.font.size = Pt(16)
            title_run.font.bold = True
        else:
            # 如果没有runs，手动添加
            title_run = title.add_run('网络设备巡检报告')
            title_run.font.size = Pt(16)
            title_run.font.bold = True
    except Exception as e:
        print(f"创建标题时出错: {e}")
        # 使用段落作为备用
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run('网络设备巡检报告')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
    
    # 基本信息
    try:
        info_heading = doc.add_heading('设备基本信息', level=1)
        if info_heading.runs:
            info_heading.runs[0].font.size = Pt(14)
    except:
        info_para = doc.add_paragraph()
        info_run = info_para.add_run('设备基本信息')
        info_run.font.size = Pt(14)
        info_run.bold = True
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # 设置表格样式
    for row in info_table.rows:
        for cell in row.cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
    
    info_table.cell(0, 0).text = '设备名称'
    info_table.cell(0, 1).text = device_info.get('name', 'N/A')
    info_table.cell(1, 0).text = '设备IP'
    info_table.cell(1, 1).text = device_info.get('ip', 'N/A')
    info_table.cell(2, 0).text = '设备厂商'
    info_table.cell(2, 1).text = device_info.get('vendor', 'N/A')
    info_table.cell(3, 0).text = '巡检时间'
    info_table.cell(3, 1).text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # 巡检结果
    try:
        results_heading = doc.add_heading('巡检结果', level=1)
        if results_heading.runs:
            results_heading.runs[0].font.size = Pt(14)
    except:
        results_para = doc.add_paragraph()
        results_run = results_para.add_run('巡检结果')
        results_run.font.size = Pt(14)
        results_run.bold = True
    
    for i, result in enumerate(inspection_results):
        try:
            # 巡检项目标题
            project_heading = doc.add_heading(f"巡检项目: {result['description']}", level=2)
            if project_heading.runs:
                project_heading.runs[0].font.size = Pt(12)
        except:
            project_para = doc.add_paragraph()
            project_run = project_para.add_run(f"巡检项目: {result['description']}")
            project_run.font.size = Pt(12)
            project_run.bold = True
        
        # 执行命令
        cmd_para = doc.add_paragraph()
        cmd_run = cmd_para.add_run("执行命令: ")
        cmd_run.bold = True
        cmd_para.add_run(result['command'])
        
        # 状态
        status_para = doc.add_paragraph()
        status_run = status_para.add_run("状态: ")
        status_run.bold = True
        status_text = status_para.add_run("成功" if result['success'] else "失败")
        if result['success']:
            status_text.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
        else:
            status_text.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        
        # 输出结果
        try:
            output_heading = doc.add_heading('输出结果:', level=3)
            if output_heading.runs:
                output_heading.runs[0].font.size = Pt(11)
        except:
            output_para = doc.add_paragraph()
            output_run = output_para.add_run('输出结果:')
            output_run.font.size = Pt(11)
            output_run.bold = True
        
        # 处理长输出，确保显示完整
        output_text = result['output']
        if len(output_text) > 10000:  # 限制输出长度
            output_text = output_text[:10000] + "\n\n...（输出过长，已截断）"
        
        # 使用等宽字体显示输出
        output_para = doc.add_paragraph()
        output_run = output_para.add_run(output_text)
        output_run.font.name = 'Consolas'
        output_run.font.size = Pt(9)
        
        # 添加分页符（除了最后一个项目）
        if i < len(inspection_results) - 1:
            doc.add_page_break()
    
    # 保存文档
    try:
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"保存Word文档时出错: {e}")
        # 尝试创建简单的文本文件作为备用
        try:
            txt_path = output_path.with_suffix('.txt')
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(f"网络设备巡检报告\n")
                f.write(f"设备名称: {device_info.get('name', 'N/A')}\n")
                f.write(f"设备IP: {device_info.get('ip', 'N/A')}\n")
                f.write(f"巡检时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                for result in inspection_results:
                    f.write(f"巡检项目: {result['description']}\n")
                    f.write(f"执行命令: {result['command']}\n")
                    f.write(f"状态: {'成功' if result['success'] else '失败'}\n")
                    f.write(f"输出结果:\n{result['output'][:5000]}\n")
                    f.write("-" * 50 + "\n")
            
            return txt_path
        except Exception as e2:
            print(f"创建文本报告也失败: {e2}")
            return None

def compare_configs(config1, config2):
    """比较两个配置文件"""
    lines1 = config1.splitlines()
    lines2 = config2.splitlines()
    
    diff = difflib.unified_diff(
        lines1, lines2,
        fromfile='config1',
        tofile='config2',
        lineterm=''
    )
    
    diff_lines = list(diff)
    
    # 分析差异
    added = []
    removed = []
    
    for line in diff_lines:
        if line.startswith('+') and not line.startswith('+++'):
            added.append(line[1:])
        elif line.startswith('-') and not line.startswith('---'):
            removed.append(line[1:])
    
    return {
        'diff': '\n'.join(diff_lines),
        'added': added,
        'removed': removed,
        'total_changes': len(added) + len(removed)
    }

def create_import_template():
    """创建设备导入模板"""
    # 创建示例数据
    data = {
        '设备名称': ['核心交换机-1', '接入交换机-1', '路由器-1'],
        'IP地址': ['192.168.1.1', '192.168.1.2', '192.168.1.3'],
        '厂商': ['cisco', 'huawei', 'h3c'],
        '用户名': ['admin', 'admin', 'admin'],
        '密码': ['password123', 'huawei@123', 'h3c@123'],
        '端口': [22, 22, 22],
        '备注': ['核心设备', '接入层设备', '边界路由器']
    }
    
    df = pd.DataFrame(data)
    
    # 保存为Excel文件
    template_path = TEMPLATE_DIR / "设备导入模板.xlsx"
    try:
        with pd.ExcelWriter(template_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='设备列表', index=False)
            
            # 添加说明工作表
            instructions = pd.DataFrame({
                '字段名': ['设备名称', 'IP地址', '厂商', '用户名', '密码', '端口', '备注'],
                '说明': [
                    '设备名称，必填',
                    '设备IP地址，必填',
                    '设备厂商（cisco/huawei/h3c/ruijie/dell/juniper/arista），必填',
                    'SSH登录用户名，必填',
                    'SSH登录密码，必填',
                    'SSH端口，默认22',
                    '设备备注信息，可选'
                ],
                '示例': [
                    '核心交换机-1',
                    '192.168.1.1',
                    'cisco',
                    'admin',
                    'password123',
                    '22',
                    '核心设备'
                ]
            })
            instructions.to_excel(writer, sheet_name='填写说明', index=False)
        
        return template_path
    except Exception as e:
        print(f"创建导入模板时出错: {e}")
        return None

def open_browser():
    """打开浏览器"""
    url = "http://localhost:8443"
    
    # 等待服务器启动
    time.sleep(2)
    
    try:
        webbrowser.open(url)
        print(f"✓ 已自动打开浏览器访问: {url}")
    except Exception as e:
        print(f"✗ 无法自动打开浏览器: {e}")
        print(f"请手动访问: {url}")

# API路由
@app.route('/')
def index():
    """主页"""
    return render_template('index.html')

@app.route('/api/devices', methods=['GET'])
def get_devices():
    """获取所有设备"""
    manager = DeviceManager()
    return jsonify({
        'success': True,
        'devices': manager.get_all_devices()
    })

@app.route('/api/devices', methods=['POST'])
def add_device():
    """添加设备"""
    data = request.json
    manager = DeviceManager()
    
    device_id = manager.add_device(
        name=data['name'],
        ip=data['ip'],
        vendor=data['vendor'],
        username=data['username'],
        password=data['password'],
        port=data.get('port', 22)
    )
    
    return jsonify({
        'success': True,
        'device_id': device_id,
        'message': '设备添加成功'
    })

@app.route('/api/devices/import', methods=['POST'])
def import_devices():
    """批量导入设备"""
    try:
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': '没有上传文件'
            })
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'message': '没有选择文件'
            })
        
        if not file.filename.endswith(('.xlsx', '.xls')):
            return jsonify({
                'success': False,
                'message': '只支持Excel文件 (.xlsx, .xls)'
            })
        
        # 读取文件内容
        excel_data = file.read()
        
        # 批量导入
        manager = DeviceManager()
        success, result = manager.batch_import(excel_data)
        
        if success:
            return jsonify({
                'success': True,
                'message': f'批量导入完成，成功: {result["success"]}，失败: {result["failed"]}',
                'data': result
            })
        else:
            return jsonify({
                'success': False,
                'message': f'导入失败: {result}'
            })
            
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'导入失败: {str(e)}'
        })

@app.route('/api/devices/template', methods=['GET'])
def download_template():
    """下载导入模板"""
    template_path = create_import_template()
    if template_path and template_path.exists():
        return send_file(template_path, as_attachment=True)
    return jsonify({'success': False, 'message': '模板文件不存在'}), 404

@app.route('/api/devices/<device_id>', methods=['DELETE'])
def delete_device(device_id):
    """删除设备"""
    manager = DeviceManager()
    success = manager.remove_device(device_id)
    
    return jsonify({
        'success': success,
        'message': '设备删除成功' if success else '设备不存在'
    })

@app.route('/api/inspection/commands', methods=['GET'])
def get_inspection_commands():
    """获取巡检命令"""
    vendor = request.args.get('vendor', 'cisco')
    commands = DEFAULT_INSPECTION_COMMANDS.get(vendor, DEFAULT_INSPECTION_COMMANDS['cisco'])
    
    return jsonify({
        'success': True,
        'commands': commands
    })

@app.route('/api/inspection/execute', methods=['POST'])
def execute_inspection():
    """执行巡检"""
    data = request.json
    device_id = data['device_id']
    commands = data['commands']
    
    manager = DeviceManager()
    device_info = manager.get_device(device_id)
    
    if not device_info:
        return jsonify({
            'success': False,
            'message': '设备不存在'
        })
    
    device = NetworkDevice(device_info)
    results = []
    
    for cmd, desc in commands:
        success, output = device.execute_command(cmd)
        results.append({
            'command': cmd,
            'description': desc,
            'success': success,
            'output': output
        })
    
    device.disconnect()
    
    # 生成报告
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f"{device_info['name']}_{device_info['ip']}_{timestamp}.docx"
    report_path = INSPECTION_DIR / report_filename
    
    try:
        report_file = create_inspection_report(device_info, results, report_path)
        
        if report_file:
            return jsonify({
                'success': True,
                'results': results,
                'report_url': f'/api/reports/download/{report_filename}'
            })
        else:
            return jsonify({
                'success': False,
                'message': '生成报告失败',
                'results': results
            })
    except Exception as e:
        print(f"生成巡检报告时出错: {e}")
        return jsonify({
            'success': False,
            'message': f'生成报告时出错: {str(e)}',
            'results': results
        })

@app.route('/api/backup', methods=['POST'])
def backup_config():
    """备份配置"""
    data = request.json
    device_id = data['device_id']
    
    manager = DeviceManager()
    device_info = manager.get_device(device_id)
    
    if not device_info:
        return jsonify({
            'success': False,
            'message': '设备不存在'
        })
    
    device = NetworkDevice(device_info)
    success, result = device.backup_config()
    device.disconnect()
    
    if success:
        return jsonify({
            'success': True,
            'message': '配置备份成功',
            'data': result
        })
    else:
        return jsonify({
            'success': False,
            'message': result
        })

@app.route('/api/backup/list', methods=['GET'])
def list_backups():
    """列出所有备份"""
    backups = []
    
    for file in CONFIG_DIR.glob("*.cfg"):
        stats = file.stat()
        backups.append({
            'filename': file.name,
            'path': str(file),
            'size': stats.st_size,
            'modified': datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'device_ip': file.name.split('_')[0]
        })
    
    return jsonify({
        'success': True,
        'backups': sorted(backups, key=lambda x: x['modified'], reverse=True)
    })

@app.route('/api/backup/package', methods=['POST'])
def package_backups():
    """打包备份文件"""
    data = request.json
    selected_files = data.get('files', [])
    
    if not selected_files:
        # 打包当天所有备份
        today = datetime.now().strftime("%Y%m%d")
        selected_files = [f.name for f in CONFIG_DIR.glob(f"*_{today}*.cfg")]
    
    if not selected_files:
        return jsonify({
            'success': False,
            'message': '没有找到备份文件'
        })
    
    # 创建ZIP文件
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_filename = f"config_backup_{timestamp}.zip"
    zip_path = CONFIG_DIR / zip_filename
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for filename in selected_files:
            filepath = CONFIG_DIR / filename
            if filepath.exists():
                zipf.write(filepath, filename)
    
    return jsonify({
        'success': True,
        'message': '打包成功',
        'zip_url': f'/api/backup/download/{zip_filename}'
    })

@app.route('/api/compare', methods=['POST'])
def compare_config():
    """比较配置文件"""
    data = request.json
    
    config1 = data.get('config1', '')
    config2 = data.get('config2', '')
    
    if not config1 or not config2:
        return jsonify({
            'success': False,
            'message': '需要两个配置文件进行比较'
        })
    
    result = compare_configs(config1, config2)
    
    return jsonify({
        'success': True,
        'result': result
    })

@app.route('/api/reports/download/<filename>')
def download_report(filename):
    """下载巡检报告"""
    # 尝试查找Word文档
    filepath = INSPECTION_DIR / filename
    
    # 如果Word文档不存在，尝试查找文本文件
    if not filepath.exists():
        txt_filename = filename.replace('.docx', '.txt')
        filepath = INSPECTION_DIR / txt_filename
    
    if filepath.exists():
        return send_file(filepath, as_attachment=True)
    return jsonify({'success': False, 'message': '文件不存在'}), 404

@app.route('/api/backup/download/<filename>')
def download_backup(filename):
    """下载备份文件"""
    filepath = CONFIG_DIR / filename
    if filepath.exists():
        return send_file(filepath, as_attachment=True)
    return jsonify({'success': False, 'message': '文件不存在'}), 404

@app.route('/api/backup/read/<filename>')
def read_backup(filename):
    """读取备份文件内容"""
    filepath = CONFIG_DIR / filename
    if filepath.exists():
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        return jsonify({
            'success': True,
            'content': content
        })
    return jsonify({'success': False, 'message': '文件不存在'}), 404

@app.route('/static/<path:filename>')
def static_files(filename):
    """提供静态文件"""
    return send_from_directory('static', filename)

def print_banner():
    """打印程序横幅"""
    banner = """
╔══════════════════════════════════════════════════════════════╗
║                                                              ║
║     网络设备巡检和配置备份软件 v1.0                         ║
║     Network Device Inspection & Backup Software             ║
║                                                              ║
╚══════════════════════════════════════════════════════════════╝
    """
    print(banner)

def main():
    """主函数"""
    # 打印程序横幅
    print_banner()
    
    print("=" * 60)
    print("程序启动中...")
    print(f"启动时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("-" * 60)
    
    # 检查必要目录
    print("检查系统目录...")
    for directory in [CONFIG_DIR, INSPECTION_DIR, TEMPLATE_DIR]:
        directory.mkdir(exist_ok=True)
        print(f"  ✓ {directory.name}: {directory}")
    
    # 创建设备导入模板
    print("创建设备导入模板...")
    template_path = create_import_template()
    if template_path:
        print(f"  ✓ 模板已创建: {template_path}")
    else:
        print("  ✗ 模板创建失败")
    
    # 加载设备数据
    print("加载设备数据...")
    manager = DeviceManager()
    device_count = len(manager.get_all_devices())
    print(f"  ✓ 已加载 {device_count} 个设备")
    
    print("-" * 60)
    print("启动Web服务器...")
    print(f"Web服务端口: 8443")
    print(f"访问地址: http://localhost:8443")
    print("-" * 60)
    
    # 在新线程中打开浏览器
    browser_thread = threading.Thread(target=open_browser, daemon=True)
    browser_thread.start()
    
    # 启动服务器
    try:
        from waitress import serve
        print("服务器已启动，按 Ctrl+C 停止服务")
        print("=" * 60)
        print("日志信息:")
        print("-" * 60)
        
        # 启动服务器
        serve(app, host='0.0.0.0', port=8443)
        
    except KeyboardInterrupt:
        print("\n" + "-" * 60)
        print("收到停止信号，正在关闭服务器...")
        print("程序已停止")
        print("=" * 60)
    except Exception as e:
        print(f"\n服务器启动失败: {e}")
        print("按任意键退出...")
        input()

if __name__ == '__main__':
    main()
